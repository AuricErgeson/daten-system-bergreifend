from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Seitenränder ────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ─── Hilfsfunktionen ─────────────────────────────────────────────

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.name = "Calibri"
    p.runs[0].font.size = Pt(11)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p

def tabelle(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "BDD7EE")
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


# ═══════════════════════════════════════════════════════════════
# TITELSEITE
# ═══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Softwaredokumentation")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Lernmaterialverwaltung")
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph()

meta = [
    ("Projekt",    "LF8 – Daten systemübergreifend bereitstellen"),
    ("Klasse",     "Anwendungsentwickler"),
    ("Abgabe",     "16.05.2026"),
    ("Datenbank",  "MySQL"),
    ("Sprache",    "Python 3"),
    ("Oberfläche", "Textbasiertes Konsolenmenü"),
]
t = doc.add_table(rows=len(meta), cols=2)
t.style = "Table Grid"
for i, (k, v) in enumerate(meta):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    t.rows[i].cells[1].text = v
for row in t.rows:
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 1. PROJEKTBESCHREIBUNG
# ═══════════════════════════════════════════════════════════════
h1("1. Projektbeschreibung")

body(
    "Die Lernmaterialverwaltung ist eine datenbankgestützte Python-Anwendung, "
    "mit der Lehrkräfte und Auszubildende digitale Lernmaterialien verwalten können. "
    "Benutzer können Dateien (z. B. PDFs, Bilder oder Python-Skripte) in das System "
    "hochladen, gezielt suchen und wieder herunterladen."
)
body(
    "Die Anwendung nutzt eine MySQL-Datenbank als zentralen Speicher. "
    "Dateien kleiner als 1 MB werden direkt in der Datenbank abgelegt; "
    "größere Dateien werden im Dateisystem gespeichert und nur als Pfad referenziert. "
    "Die Bedienung erfolgt über ein einfaches Textmenü in der Konsole – ohne grafische Oberfläche."
)
body(
    "Jedes Lernmaterial gehört zu einem Themengebiet (z. B. Informatik oder Mathematik) "
    "und kann mit Kategorien, Tags und Kommentaren versehen werden. "
    "Revisionssicherheit wird durch automatische Zeitstempel und eine Versionstabelle gewährleistet."
)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# 2. ZIELSETZUNG
# ═══════════════════════════════════════════════════════════════
h1("2. Zielsetzung")

body(
    "Ziel des Projekts ist es, die bisher unstrukturierte Ablage von Lernmaterialien "
    "in verschiedenen Verzeichnissen durch eine zentrale, durchsuchbare Datenbank zu ersetzen. "
    "Dadurch entfällt das manuelle Suchen in Ordnern und das Risiko, veraltete Dateiversionen weiterzugeben."
)
body(
    "Sowohl Lehrkräfte als auch Auszubildende sollen Materialien hochladen, suchen und herunterladen können. "
    "Metadaten wie Autor, Erstelldatum, Dateityp und Themengebiet machen jede Datei eindeutig auffindbar. "
    "Kommentare ermöglichen Rückmeldungen direkt am jeweiligen Material."
)
body(
    "Die Anwendung soll außerdem zeigen, wie Daten systemübergreifend bereitgestellt werden können – "
    "ein zentrales Thema im Lernfeld 8 der Berufsausbildung zum Anwendungsentwickler."
)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# 3. FUNKTIONSÜBERSICHT
# ═══════════════════════════════════════════════════════════════
h1("3. Funktionsübersicht")

body("Das Hauptmenü der Anwendung bietet folgende Funktionen:")
doc.add_paragraph()

tabelle(
    ["Menüpunkt", "Funktion", "Beschreibung"],
    [
        ["1", "Material hochladen",
         "Benutzer wählt eine Datei aus, gibt Titel, Themengebiet, Autor und Kategorie an. "
         "Dateien unter 1 MB werden in die Datenbank gespeichert, größere Dateien ins Dateisystem."],
        ["2", "Material herunterladen",
         "Zeigt alle vorhandenen Materialien als Tabelle an. "
         "Benutzer wählt eine ID und einen Zielordner – die Datei wird dort gespeichert."],
        ["3", "Materialien suchen",
         "Suche nach Titel (Texteingabe), Themengebiet oder Autor. "
         "Filter können kombiniert werden. Ergebnisse werden tabellarisch angezeigt."],
        ["4", "Material löschen",
         "Zeigt alle Materialien an. Benutzer wählt eine ID und bestätigt das Löschen. "
         "Dateien im Dateisystem werden dabei ebenfalls entfernt."],
        ["5", "Kommentare verwalten",
         "Kommentare zu einem Material anzeigen, neue Kommentare hinzufügen "
         "oder vorhandene Kommentare löschen."],
        ["6", "Listen & Übersichten",
         "Alle Einträge der Haupttabellen anzeigen: Materialien, Benutzer, "
         "Themengebiete, Kategorien und Tags."],
        ["7", "Datenbank-Abfragen",
         "Sieben vorgefertigte SQL-Abfragen ausführen (COUNT, AVG, JOIN, Multi-JOIN). "
         "Alle 7 Abfragen können auch auf einmal gestartet werden."],
        ["8", "Neue Einträge erstellen",
         "Benutzer, Themengebiete, Kategorien und Tags anlegen. "
         "Diese werden dann beim Hochladen von Materialien verwendet."],
    ],
    col_widths=[1.2, 3.8, 10.0],
)


# ═══════════════════════════════════════════════════════════════
# 4. IMPLEMENTIERUNGSDETAILS
# ═══════════════════════════════════════════════════════════════
h1("4. Implementierungsdetails")

h2("4.1 Verwendete Python-Bibliotheken")

tabelle(
    ["Bibliothek", "Zweck"],
    [
        ["SQLAlchemy",    "Verbindung zur MySQL-Datenbank und Zugriff auf Tabellen über Python-Objekte (ORM). "
                          "Die Tabellen werden mit automap_base automatisch aus der Datenbank gelesen – "
                          "es müssen keine Klassen manuell geschrieben werden."],
        ["PyMySQL",       "MySQL-Treiber für Python. SQLAlchemy nutzt ihn intern, "
                          "um tatsächlich mit der Datenbank zu kommunizieren."],
        ["python-dotenv", "Lädt Zugangsdaten (Datenbankname, Passwort, Host) aus einer .env-Datei. "
                          "Passwörter stehen so nicht direkt im Quellcode."],
        ["rich",          "Gibt Tabellen und farbige Meldungen in der Konsole aus. "
                          "Macht das Textmenü übersichtlicher und leichter lesbar."],
        ["os / shutil",   "Standardbibliotheken für Dateizugriff: Pfade prüfen, "
                          "Dateien kopieren und Ordner erstellen."],
    ],
    col_widths=[4.0, 11.0],
)

h2("4.2 Aufbau des Codes")

body(
    "Der Code ist in drei Dateien aufgeteilt, jede mit einer klaren Aufgabe:"
)

tabelle(
    ["Datei", "Inhalt"],
    [
        ["database.py",
         "Stellt die Datenbankverbindung her. Lädt Zugangsdaten aus der .env-Datei, "
         "baut die MySQL-URL zusammen und liest alle Tabellen automatisch ein (automap). "
         "Jede Tabelle (z. B. material, benutzer) wird als Python-Klasse verfügbar gemacht."],
        ["services.py",
         "Enthält alle Datenbankfunktionen: Materialien hochladen, herunterladen, suchen, löschen. "
         "Außerdem alle 7 Abfragefunktionen. Diese Datei kennt keine Benutzereingaben – "
         "sie gibt nur Daten zurück oder speichert sie."],
        ["main.py",
         "Die Benutzerschnittstelle. Zeigt das Menü an, liest Eingaben ein und "
         "ruft die passenden Funktionen aus services.py auf. "
         "Gibt Ergebnisse formatiert mit rich in der Konsole aus."],
    ],
    col_widths=[3.5, 11.5],
)

h2("4.3 Speicherstrategie")

body(
    "Die Anwendung unterscheidet bei jedem Upload automatisch zwischen zwei Speicherwegen:"
)

tabelle(
    ["Dateigröße", "Speicherweg", "Feld in der Datenbank"],
    [
        ["Kleiner als 1 MB  (< 1.048.576 Bytes)",
         "Dateiinhalt direkt in die Datenbank (LONGBLOB-Spalte inhalt)",
         "is_in_database = TRUE, speicherort = NULL"],
        ["1 MB oder größer",
         "Datei wird in den Ordner uploads/ kopiert. "
         "In der Datenbank wird nur der Dateipfad gespeichert.",
         "is_in_database = FALSE, speicherort = Pfad zur Datei"],
    ],
    col_widths=[4.5, 7.0, 3.5],
)

body(
    "Beim Download prüft die Anwendung das Feld is_in_database und liest die Datei "
    "entweder aus der Datenbank oder vom Dateisystem – für den Benutzer transparent."
)

h2("4.4 Die 7 Datenbankabfragen")

body(
    "Die Anwendung enthält sieben vorgefertigte Abfragen, die verschiedene SQL-Techniken zeigen:"
)

tabelle(
    ["Nr.", "Abfragename", "SQL-Technik", "Was sie liefert"],
    [
        ["1", "Materialien pro Themengebiet",
         "COUNT + GROUP BY (Aggregation)",
         "Wie viele Materialien es pro Themengebiet gibt."],
        ["2", "Ø Dateigröße pro Dateityp",
         "AVG + GROUP BY (Aggregation)",
         "Durchschnittliche Dateigröße für jeden Dateityp (PDF, DOCX …)."],
        ["3", "Materialien mit Themengebiet",
         "INNER JOIN (material → themengebiet)",
         "Titel und Dateiname jedes Materials mit dem Namen seines Themengebiets."],
        ["4", "Materialien mit Kommentaren",
         "INNER JOIN (material → kommentar)",
         "Alle Materialien, die mindestens einen Kommentar haben, mit Kommentartext."],
        ["5", "Kommentare pro Material",
         "INNER JOIN + COUNT (Join + Aggregation)",
         "Für jedes Material: wie viele Kommentare sind vorhanden."],
        ["6", "Material + Autor + Thema",
         "2× INNER JOIN (material → benutzer + themengebiet)",
         "Titel, Autor-Name und Themengebiet jedes Materials in einer Zeile."],
        ["7", "Vollständige Suche mit Filtern",
         "3× INNER JOIN (material → benutzer + themengebiet + kategorie)",
         "Alle Details eines Materials: Titel, Typ, Autor, Thema und Kategorie. "
         "Filterbar nach Thema und Autor."],
    ],
    col_widths=[0.8, 4.5, 4.2, 5.5],
)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# 5. SQL-CODE
# ═══════════════════════════════════════════════════════════════
h1("5. SQL-Code (schema.sql)")

body("Das folgende SQL-Skript erstellt die komplette Datenbank mit allen Tabellen und Triggern:")
doc.add_paragraph()

sql_code = r"""DROP DATABASE IF EXISTS lernmaterialverwaltung;

CREATE DATABASE lernmaterialverwaltung
    CHARACTER SET utf8mb4
    COLLATE utf8mb4_unicode_ci;

USE lernmaterialverwaltung;

CREATE TABLE themengebiet (
    id           INT AUTO_INCREMENT PRIMARY KEY,
    name         VARCHAR(100) NOT NULL UNIQUE,
    beschreibung TEXT
);

CREATE TABLE benutzer (
    id         INT AUTO_INCREMENT PRIMARY KEY,
    username   VARCHAR(100) NOT NULL UNIQUE,
    email      VARCHAR(255) NOT NULL UNIQUE,
    rolle      ENUM('Lehrer', 'Azubi', 'Admin') NOT NULL DEFAULT 'Azubi',
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);

CREATE TABLE kategorie (
    id           INT AUTO_INCREMENT PRIMARY KEY,
    name         VARCHAR(100) NOT NULL UNIQUE,
    beschreibung TEXT
);

CREATE TABLE tag (
    id   INT AUTO_INCREMENT PRIMARY KEY,
    name VARCHAR(50) NOT NULL UNIQUE
);

CREATE TABLE material (
    id               INT AUTO_INCREMENT PRIMARY KEY,
    titel            VARCHAR(255) NOT NULL,
    beschreibung     TEXT,
    themengebiet_id  INT NOT NULL,
    benutzer_id      INT NOT NULL,
    kategorie_id     INT NOT NULL,
    erstellungsdatum TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    aenderungsdatum  TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                     ON UPDATE CURRENT_TIMESTAMP,
    dateiname        VARCHAR(255),
    dateityp         VARCHAR(50),
    dateigroesse     BIGINT DEFAULT 0,
    speicherort      VARCHAR(500),
    inhalt           LONGBLOB,
    is_in_database   BOOLEAN DEFAULT FALSE,

    FOREIGN KEY (themengebiet_id) REFERENCES themengebiet(id) ON DELETE RESTRICT,
    FOREIGN KEY (benutzer_id)     REFERENCES benutzer(id)     ON DELETE RESTRICT,
    FOREIGN KEY (kategorie_id)    REFERENCES kategorie(id)    ON DELETE RESTRICT
);

CREATE TABLE kommentar (
    id               INT AUTO_INCREMENT PRIMARY KEY,
    material_id      INT NOT NULL,
    autor_id         INT NOT NULL,
    text             TEXT NOT NULL,
    erstellungsdatum TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    aenderungsdatum  TIMESTAMP DEFAULT CURRENT_TIMESTAMP,

    FOREIGN KEY (material_id) REFERENCES material(id)  ON DELETE CASCADE,
    FOREIGN KEY (autor_id)    REFERENCES benutzer(id)  ON DELETE RESTRICT
);

CREATE TABLE material_tag (
    id          INT AUTO_INCREMENT PRIMARY KEY,
    material_id INT NOT NULL,
    tag_id      INT NOT NULL,

    FOREIGN KEY (material_id) REFERENCES material(id) ON DELETE CASCADE,
    FOREIGN KEY (tag_id)      REFERENCES tag(id)      ON DELETE CASCADE
);

CREATE TABLE version (
    id              INT AUTO_INCREMENT PRIMARY KEY,
    material_id     INT NOT NULL,
    autor_id        INT NOT NULL,
    version_nr      INT NOT NULL,
    aenderungsdatum TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    aenderungen     TEXT,

    FOREIGN KEY (material_id) REFERENCES material(id) ON DELETE CASCADE,
    FOREIGN KEY (autor_id)    REFERENCES benutzer(id) ON DELETE RESTRICT
);

DELIMITER //

CREATE TRIGGER before_material_update
BEFORE UPDATE ON material
FOR EACH ROW
BEGIN
    SET NEW.aenderungsdatum = CURRENT_TIMESTAMP;
END //

CREATE TRIGGER before_kommentar_update
BEFORE UPDATE ON kommentar
FOR EACH ROW
BEGIN
    SET NEW.aenderungsdatum = CURRENT_TIMESTAMP;
END //

DELIMITER ;"""

code_block(sql_code)


# ═══════════════════════════════════════════════════════════════
# 6. ERD UND ERM
# ═══════════════════════════════════════════════════════════════
h1("6. ERD und ERM – Entitäten und Beziehungen")

body(
    "Da in diesem Dokument keine Grafiken erstellt werden können, wird das "
    "Entity-Relationship-Diagramm (ERD) hier in Textform beschrieben. "
    "Das ERD kann anschließend mit draw.io (kostenlos im Browser unter app.diagrams.net) "
    "gezeichnet werden."
)
doc.add_paragraph()

entities = [
    (
        "benutzer",
        "Speichert alle Benutzer des Systems.",
        "id (PK), username, email, rolle (Lehrer / Azubi / Admin), created_at",
        [
            "Kann viele Materialien hochladen (1:N → material.benutzer_id)",
            "Kann viele Kommentare schreiben (1:N → kommentar.autor_id)",
            "Kann viele Versionen erstellen (1:N → version.autor_id)",
        ],
    ),
    (
        "themengebiet",
        "Ordnet Materialien einem Fachbereich zu (z. B. Informatik, Mathematik).",
        "id (PK), name, beschreibung",
        [
            "Kann viele Materialien enthalten (1:N → material.themengebiet_id)",
        ],
    ),
    (
        "kategorie",
        "Weitere Klassifizierung von Materialien (z. B. Übungsaufgabe, Skript).",
        "id (PK), name, beschreibung",
        [
            "Kann viele Materialien haben (1:N → material.kategorie_id)",
        ],
    ),
    (
        "tag",
        "Schlagwörter für schnelle Suche und Filterung.",
        "id (PK), name",
        [
            "Viele Tags können zu vielen Materialien gehören (N:M über material_tag)",
        ],
    ),
    (
        "material",
        "Zentrales Element: speichert alle Lernmaterialien mit Metadaten und Dateiinhalt.",
        "id (PK), titel, beschreibung, themengebiet_id (FK), benutzer_id (FK), "
        "kategorie_id (FK), erstellungsdatum, aenderungsdatum, dateiname, dateityp, "
        "dateigroesse, speicherort, inhalt (LONGBLOB), is_in_database",
        [
            "Gehört zu genau 1 themengebiet (N:1)",
            "Gehört zu genau 1 benutzer als Autor (N:1)",
            "Gehört zu genau 1 kategorie (N:1)",
            "Kann viele kommentare haben (1:N)",
            "Kann viele tags haben (N:M über material_tag)",
            "Kann viele versionen haben (1:N)",
        ],
    ),
    (
        "kommentar",
        "Rückmeldungen und Hinweise zu einem Material.",
        "id (PK), material_id (FK), autor_id (FK), text, erstellungsdatum, aenderungsdatum",
        [
            "Gehört zu genau 1 material (N:1)",
            "Wurde geschrieben von genau 1 benutzer (N:1)",
        ],
    ),
    (
        "material_tag",
        "Verbindungstabelle für die N:M-Beziehung zwischen Material und Tag.",
        "id (PK), material_id (FK), tag_id (FK)",
        [
            "Verbindet material und tag (Auflösung der N:M-Beziehung)",
        ],
    ),
    (
        "version",
        "Speichert Änderungshistorie zu einem Material (Versionsverwaltung).",
        "id (PK), material_id (FK), autor_id (FK), version_nr, aenderungsdatum, aenderungen",
        [
            "Gehört zu genau 1 material (N:1)",
            "Wurde erstellt von genau 1 benutzer (N:1)",
        ],
    ),
]

for name, desc, attrs, rels in entities:
    h2(f"Entität: {name}")
    body(desc)
    p = doc.add_paragraph()
    run = p.add_run("Attribute: ")
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(attrs).font.size = Pt(11)
    p = doc.add_paragraph()
    run = p.add_run("Beziehungen:")
    run.bold = True
    run.font.size = Pt(11)
    for rel in rels:
        rp = doc.add_paragraph(rel, style="List Bullet")
        rp.runs[0].font.size = Pt(11)
    doc.add_paragraph()

body(
    "Hinweis: Das ERD ist nach dem Lesen dieser Beschreibung in draw.io zu zeichnen. "
    "Jede Entität wird als Rechteck dargestellt. "
    "Beziehungen werden mit Linien und Kardinalitätsangaben (1, N, M) verbunden."
)

doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# 7. NORMALISIERTE TABELLENSTRUKTUR
# ═══════════════════════════════════════════════════════════════
h1("7. Normalisierte Tabellenstruktur")

body(
    "Alle Tabellen liegen in der dritten Normalform (3NF) vor: "
    "Es gibt keine Wiederholungsgruppen (1NF), alle Spalten hängen vollständig vom Primärschlüssel ab (2NF) "
    "und keine Spalte hängt von einer Nicht-Schlüsselspalte ab (3NF)."
)
doc.add_paragraph()

tables_data = [
    ("benutzer", [
        ("id",         "INT",          "PK, AUTO_INCREMENT",  "Eindeutige Benutzer-ID"),
        ("username",   "VARCHAR(100)", "NOT NULL, UNIQUE",    "Benutzername (Login-Name)"),
        ("email",      "VARCHAR(255)", "NOT NULL, UNIQUE",    "E-Mail-Adresse"),
        ("rolle",      "ENUM",         "NOT NULL",            "Benutzerrolle: Lehrer, Azubi oder Admin"),
        ("created_at", "TIMESTAMP",    "DEFAULT now()",       "Zeitpunkt der Kontoerstellung"),
    ]),
    ("themengebiet", [
        ("id",           "INT",          "PK, AUTO_INCREMENT", "Eindeutige ID"),
        ("name",         "VARCHAR(100)", "NOT NULL, UNIQUE",   "Name des Themengebiets (z. B. Informatik)"),
        ("beschreibung", "TEXT",         "optional",           "Kurze Erklärung zum Themengebiet"),
    ]),
    ("kategorie", [
        ("id",           "INT",          "PK, AUTO_INCREMENT", "Eindeutige ID"),
        ("name",         "VARCHAR(100)", "NOT NULL, UNIQUE",   "Name der Kategorie (z. B. Übungsaufgabe)"),
        ("beschreibung", "TEXT",         "optional",           "Kurze Erklärung zur Kategorie"),
    ]),
    ("tag", [
        ("id",   "INT",         "PK, AUTO_INCREMENT", "Eindeutige Tag-ID"),
        ("name", "VARCHAR(50)", "NOT NULL, UNIQUE",   "Schlagwort (z. B. Python, SQL)"),
    ]),
    ("material", [
        ("id",               "INT",          "PK, AUTO_INCREMENT",   "Eindeutige Material-ID"),
        ("titel",            "VARCHAR(255)", "NOT NULL",             "Titel des Lernmaterials"),
        ("beschreibung",     "TEXT",         "optional",             "Kurzbeschreibung des Inhalts"),
        ("themengebiet_id",  "INT",          "FK → themengebiet.id", "Zugehöriges Themengebiet"),
        ("benutzer_id",      "INT",          "FK → benutzer.id",     "Autor / Hochladender Benutzer"),
        ("kategorie_id",     "INT",          "FK → kategorie.id",    "Zugehörige Kategorie"),
        ("erstellungsdatum", "TIMESTAMP",    "DEFAULT now()",        "Datum und Uhrzeit des Uploads"),
        ("aenderungsdatum",  "TIMESTAMP",    "AUTO ON UPDATE",       "Letzte Änderung (automatisch)"),
        ("dateiname",        "VARCHAR(255)", "optional",             "Originaler Dateiname"),
        ("dateityp",         "VARCHAR(50)",  "optional",             "MIME-Typ (z. B. application/pdf)"),
        ("dateigroesse",     "BIGINT",       "DEFAULT 0",            "Größe der Datei in Bytes"),
        ("speicherort",      "VARCHAR(500)", "optional",             "Pfad im Dateisystem (nur bei großen Dateien)"),
        ("inhalt",           "LONGBLOB",     "optional",             "Dateiinhalt (nur bei Dateien < 1 MB)"),
        ("is_in_database",   "BOOLEAN",      "DEFAULT FALSE",        "TRUE = Inhalt in DB; FALSE = Pfad im Dateisystem"),
    ]),
    ("kommentar", [
        ("id",               "INT",       "PK, AUTO_INCREMENT",  "Eindeutige Kommentar-ID"),
        ("material_id",      "INT",       "FK → material.id",    "Zugehöriges Lernmaterial"),
        ("autor_id",         "INT",       "FK → benutzer.id",    "Verfasser des Kommentars"),
        ("text",             "TEXT",      "NOT NULL",            "Kommentartext"),
        ("erstellungsdatum", "TIMESTAMP", "DEFAULT now()",       "Datum der Ersterstellung"),
        ("aenderungsdatum",  "TIMESTAMP", "DEFAULT now()",       "Datum der letzten Änderung"),
    ]),
    ("material_tag", [
        ("id",          "INT", "PK, AUTO_INCREMENT", "Eindeutige ID der Verknüpfung"),
        ("material_id", "INT", "FK → material.id",   "Verweis auf das Material"),
        ("tag_id",      "INT", "FK → tag.id",         "Verweis auf den Tag"),
    ]),
    ("version", [
        ("id",              "INT",       "PK, AUTO_INCREMENT", "Eindeutige Versions-ID"),
        ("material_id",     "INT",       "FK → material.id",   "Zugehöriges Material"),
        ("autor_id",        "INT",       "FK → benutzer.id",   "Benutzer, der die Änderung vorgenommen hat"),
        ("version_nr",      "INT",       "NOT NULL",           "Versionsnummer (1, 2, 3 …)"),
        ("aenderungsdatum", "TIMESTAMP", "DEFAULT now()",      "Zeitpunkt der Änderung"),
        ("aenderungen",     "TEXT",      "optional",           "Beschreibung der vorgenommenen Änderungen"),
    ]),
]

for tname, cols in tables_data:
    h2(f"Tabelle: {tname}")
    tabelle(
        ["Spalte", "Datentyp", "Einschränkung", "Beschreibung"],
        cols,
        col_widths=[3.5, 3.0, 3.5, 5.0],
    )


# ─── Speichern ──────────────────────────────────────────────────
output_path = "LF8_Dokumentation.docx"
doc.save(output_path)
print(f"Dokumentation gespeichert: {output_path}")
